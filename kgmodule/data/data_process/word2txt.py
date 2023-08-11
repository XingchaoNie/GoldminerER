import os.path

import docx


def docx_to_txt():
    # 打开文件
    file = docx.Document('../../data/data_process/unlabel/unlabel0.docx')
    # file = docx.Document('../../data/data_process/unlabel/test.ocx')
    with open('../../data/data_process/unlabel/unlabel0.txt', 'w', encoding='utf-8') as f:
        # print(file.paragraphs)
        for para in file.paragraphs:
            # print(para)
            if len(para.text) != 0:
                cnt = para.text.replace(' ', '')
                f.write(cnt)
                # print(cnt)
                f.write("\n")            # f.writelines(para.text)
    f.close()

# 调用函数
if __name__ == '__main__':
    docx_to_txt()
    # father = os.path.dirname
    # filepath = os.path.join(father(father(father(father(os.path.abspath(__file__))))), r'/files/txt/')
    # print(filepath)

